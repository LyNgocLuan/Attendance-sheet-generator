import os
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt

from constants import *
from calendar import Calendar

from time import time


def validate_month(month):
    while True:
        if month > 12 or month < 1:
            print("Loi: Nhap tu thang 1 toi thang 12")
            continue
        return month


def get_days_in_week(year, month):
    calendar_month_days = Calendar().monthdays2calendar(year, month)
    days_in_week = list(
        map(lambda x: [d[1] for d in x if d[0] != 0], (month_days for month_days in calendar_month_days)))
    week_days = []
    while days_in_week:
        week_days.extend(days_in_week.pop(0))
    return week_days


def set_title(doc, month, year):
    title = d.paragraphs[0].runs[0]

    replace_start = title.text.find("(")
    replace_end = title.text.find(")")

    to_be_replaced = title.text[replace_start + 1:replace_end]
    to_replace = "Tháng {0}/{1}".format(month, year)

    new_title = title
    new_title.text = title.text.replace(to_be_replaced, to_replace)
    new_title.font.name = "Times New Roman"

    doc.paragraphs[0].runs[0] = new_title


def set_month_week_days(table, days_in_month, days_in_week):
    i = 0
    while i < len(days_in_month):
        table.cell(DOM_ROW, i + 2).text = str(days_in_month[i])
        table.cell(DOM_ROW, i + 2).paragraphs[0].runs[0].font.bold = True
        table.cell(DOM_ROW, i + 2).paragraphs[0].runs[0].font.name = "Times New Roman"
        table.cell(DOM_ROW, i + 2).paragraphs[0].runs[0].font.size = Pt(10)

        table.cell(DOW_ROW, i + 2).text = DAY_TO_TEXT[days_in_week[i]]
        table.cell(DOW_ROW, i + 2).paragraphs[0].runs[0].font.bold = True
        table.cell(DOW_ROW, i + 2).paragraphs[0].runs[0].font.name = "Times New Roman"
        table.cell(DOW_ROW, i + 2).paragraphs[0].runs[0].font.size = Pt(10)

        i += 1
    for i in range(2):
        for cell in table.rows[i].cells[len(days_in_month) + 2:]:
            cell.text = ''


def set_sunday_column_color(table, color):
    for i in range(MAX_MONTH_DAYS):
        if table.cell(DOW_ROW, i + 2).text == "CN":
            for j in range(2, len(table.rows)):
                set_cell_color(color, table.cell(j, i + 2))


def set_cell_color(color, cell):
    shading_elm = parse_xml(r'<w:shd {0} w:fill="{1}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)


if __name__ == '__main__':
    start = time()
    path = Path()
    docs = [p for p in path.glob("*.docx")]

    while True:
        try:
            month = validate_month(int(input("Nhap thang: ")))
            year = int(input("Nhap nam: "))
            break
        except ValueError:
            print("Loi: Nhap thang va nam bang so")
            continue

    new_path = Path("Tháng {0}-{1}".format(month, year))
    if not new_path.exists():
        new_path.mkdir(parents=True)

    days_in_month = [d for d in Calendar().itermonthdays(year, month) if d != 0]
    days_in_week = get_days_in_week(year, month)

    print("Running...")

    for doc in docs:
        d = Document(doc.name)

        set_title(d, month, year)

        table = d.tables[0]

        set_sunday_column_color(table, "FFFFFF")

        set_month_week_days(table, days_in_month, days_in_week)

        set_sunday_column_color(table, "696969")

        d.save(str(new_path.absolute()) + "\\" + doc.name)

        print(doc.name + "...Done")

    print("Execution time: ", time() - start)
    print("Done")
    os.system("pause")
